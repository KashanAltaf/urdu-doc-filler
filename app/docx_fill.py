from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docxtpl import DocxTemplate
from lxml import etree

URDU_FONT = "Jameel Noori Nastaleeq"
PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def extract_placeholders(docx_bytes: bytes) -> list[str]:
    """Find unique {{placeholder}} names in a .docx template."""
    doc = Document(BytesIO(docx_bytes))
    found: list[str] = []
    seen: set[str] = set()

    def consider(text: str) -> None:
        for match in PLACEHOLDER_RE.finditer(text or ""):
            name = match.group(1).strip()
            if name and name not in seen:
                seen.add(name)
                found.append(name)

    for para in doc.paragraphs:
        consider(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    consider(para.text)
    # Do NOT touch section.header / footer — python-docx creates empty headers.

    return found


def _set_run_font(run, font_name: str = URDU_FONT, size_pt: float | None = None) -> None:
    """Set font family only; never alter existing font size unless size_pt is given."""
    if run._element.find(qn("w:drawing")) is not None:
        return
    if not (run.text or "").strip():
        return
    # Preserve current size before any font API side-effects
    existing_size = run.font.size
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    elif existing_size is not None and run.font.size != existing_size:
        run.font.size = existing_size



def _ensure_para_rtl(paragraph) -> None:
    """Set Word 'Right-to-Left Text Direction' (w:bidi) on a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))


def _ensure_para_align_right(paragraph) -> None:
    """Force Align Right so Word's Align Right control is active."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "right")
    # Drop conflicting indentation that can make RTL text look left-stuck
    for tag in ("w:ind", "w:framePr"):
        el = pPr.find(qn(tag))
        if el is not None:
            pPr.remove(el)


def _ensure_table_rtl(table) -> None:
    """Mark table as RTL visual so Urdu columns/alignment behave correctly."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    if tblPr.find(qn("w:bidiVisual")) is None:
        # place near start of tblPr
        bidi = OxmlElement("w:bidiVisual")
        tblPr.insert(0, bidi)


def _ensure_run_rtl(run) -> None:
    """Mark a text run as RTL complex-script."""
    if run._element.find(qn("w:drawing")) is not None:
        return
    if not (run.text or "").strip():
        return
    rPr = run._element.get_or_add_rPr()
    if rPr.find(qn("w:rtl")) is None:
        rPr.append(OxmlElement("w:rtl"))


def apply_urdu_font(docx_path: Path, font_name: str = URDU_FONT) -> None:
    """Force Urdu font, RTL, and right-align (except school name / logo)."""
    doc = Document(str(docx_path))

    # First body paragraph = school name + logo — keep centered
    def style_paragraphs(paragraphs, *, force_right: bool) -> None:
        for para in paragraphs:
            _ensure_para_rtl(para)
            if force_right:
                _ensure_para_align_right(para)
            for run in para.runs:
                if run._element.find(qn("w:drawing")) is not None:
                    continue
                if not (run.text or "").strip():
                    continue
                _set_run_font(run, font_name)
                _ensure_run_rtl(run)

    for idx, para in enumerate(doc.paragraphs):
        if idx == 0:
            # Logo + school name: keep center alignment
            _ensure_para_rtl(para)
            pPr = para._p.get_or_add_pPr()
            jc = pPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                pPr.append(jc)
            jc.set(qn("w:val"), "center")
            for run in para.runs:
                if run._element.find(qn("w:drawing")) is not None:
                    continue
                if not (run.text or "").strip():
                    continue
                _set_run_font(run, font_name)
                _ensure_run_rtl(run)
        else:
            style_paragraphs([para], force_right=True)

    for table in doc.tables:
        _ensure_table_rtl(table)
        for row in table.rows:
            for cell in row.cells:
                style_paragraphs(cell.paragraphs, force_right=True)


    for section in doc.sections:
        sectPr = section._sectPr
        if sectPr is not None and sectPr.find(qn("w:bidi")) is None:
            sectPr.insert(0, OxmlElement("w:bidi"))

    doc.save(str(docx_path))
    restore_original_headers(docx_path)



def restore_original_headers(docx_path: Path) -> None:
    """
    Keep only the original even-page header (header1).
    Drop any default/empty headers introduced by python-docx saves.
    """
    path = Path(docx_path)
    tmp = path.with_suffix(".tmp.docx")

    with ZipFile(path, "r") as zin:
        namelist = zin.namelist()
        if "word/header1.xml" not in namelist:
            return

        rels_root = etree.fromstring(zin.read("word/_rels/document.xml.rels"))
        logo_rid = None
        for rel in rels_root:
            if rel.get("Target", "").endswith("header1.xml"):
                logo_rid = rel.get("Id")
                break
        if not logo_rid:
            return

        doc_root = etree.fromstring(zin.read("word/document.xml"))
        ns = {"w": W_NS}
        for sect in doc_root.findall(".//w:sectPr", ns):
            for href in list(sect.findall("w:headerReference", ns)):
                sect.remove(href)
            for fref in list(sect.findall("w:footerReference", ns)):
                sect.remove(fref)
            # Original file only had even header
            el = etree.Element(f"{{{W_NS}}}headerReference")
            el.set(f"{{{W_NS}}}type", "even")
            el.set(f"{{{R_NS}}}id", logo_rid)
            sect.insert(0, el)

        new_doc = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        with ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
            for item in namelist:
                if item == "word/document.xml":
                    zout.writestr(item, new_doc)
                elif item == "word/_rels/document.xml.rels":
                    continue
                elif item in ("word/header2.xml", "word/footer1.xml"):
                    continue
                else:
                    zout.writestr(item, zin.read(item))
            for rel in list(rels_root):
                target = rel.get("Target", "")
                if target in ("header2.xml", "footer1.xml"):
                    rels_root.remove(rel)
            zout.writestr(
                "word/_rels/document.xml.rels",
                etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True),
            )

    tmp.replace(path)


def fill_template(
    template_bytes: bytes,
    context: dict[str, Any],
    output_path: Path,
    font_name: str = URDU_FONT,
) -> Path:
    """Render placeholders with docxtpl, then clean empties and apply font."""
    tpl = DocxTemplate(BytesIO(template_bytes))
    safe_context: dict[str, Any] = {}
    for key, value in context.items():
        if isinstance(value, (list, dict, bool, int, float)):
            safe_context[str(key)] = value
        elif value is None:
            safe_context[str(key)] = ""
        else:
            safe_context[str(key)] = str(value)
    safe_context.setdefault("outcomes", [])
    safe_context.setdefault("success_items", [])
    safe_context.setdefault("activities", [])
    tpl.render(safe_context)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tpl.save(str(output_path))
    _cleanup_empty_content(output_path)
    apply_urdu_font(output_path, font_name=font_name)
    return output_path


def _cleanup_empty_content(docx_path: Path) -> None:
    """Remove blank paragraphs in objective cells and empty activity rows."""
    doc = Document(str(docx_path))
    if len(doc.tables) >= 3:
        t2 = doc.tables[2]
        for ci in (0, 2):
            cell = t2.rows[0].cells[ci]
            for p in list(cell.paragraphs):
                if not (p.text or "").strip() and len(cell.paragraphs) > 1:
                    p._element.getparent().remove(p._element)
    if len(doc.tables) >= 4:
        t3 = doc.tables[3]
        for ri in range(len(t3.rows) - 1, 0, -1):
            texts = [(c.text or "").strip() for c in t3.rows[ri].cells]
            if not any(texts):
                t3._tbl.remove(t3.rows[ri]._tr)
    doc.save(str(docx_path))
    restore_original_headers(docx_path)
