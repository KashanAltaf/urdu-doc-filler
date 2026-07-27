"""Extract plain text from uploaded PDF or DOCX books."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    raise ValueError("صرف PDF یا DOCX فائلیں قبول ہیں۔")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(f"[صفحہ {i}]\n{text}")
    return "\n\n".join(parts).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def chunk_text(text: str, chunk_size: int = 900, overlap: int = 150) -> list[str]:
    """Simple character-based chunks with overlap (works for Urdu)."""
    text = (text or "").strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_size, n)
        # prefer break on newline/space near the end
        if end < n:
            window = text[start:end]
            br = max(window.rfind("\n"), window.rfind("۔"), window.rfind(" "))
            if br > chunk_size // 3:
                end = start + br + 1
        piece = text[start:end].strip()
        if piece:
            chunks.append(piece)
        if end >= n:
            break
        start = max(0, end - overlap)
    return chunks
