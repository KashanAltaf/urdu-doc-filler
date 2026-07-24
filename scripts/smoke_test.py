from pathlib import Path
import json
from app.docx_fill import fill_template, extract_placeholders

ROOT = Path(__file__).resolve().parent.parent
tpl = ROOT / "templates" / "lesson_plan_template.docx"
defaults = json.loads((ROOT / "app" / "default_values.json").read_text(encoding="utf-8"))
# tweak one field to prove mapping
defaults = dict(defaults)
defaults["date"] = "۲۵ اگست"
defaults["school_name"] = "دی لنکس اسکول (ٹیسٹ)"
out = ROOT / "outputs" / "smoke-test.docx"
fill_template(tpl.read_bytes(), defaults, out)
# verify no leftover placeholders and values present
from docx import Document
doc = Document(str(out))
text = "\n".join(p.text for p in doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            text += "\n" + cell.text
left = extract_placeholders(out.read_bytes())
ok_school = "دی لنکس اسکول (ٹیسٹ)" in text
ok_date = "۲۵ اگست" in text
print("leftover_placeholders", left)
print("school_ok", ok_school)
print("date_ok", ok_date)
print("out", out)
assert not left
assert ok_school and ok_date
print("SMOKE_OK")
