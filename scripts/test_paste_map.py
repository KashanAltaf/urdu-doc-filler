from pathlib import Path
import json
from app.docx_fill import fill_template
from app.lesson_parse import parse_lesson_paste, context_from_flat_fields

SAMPLE = """سبق کا منصوبہ (جماعت سوم)

ہفتہ: دوسرا
سنگل پیریڈ: 40 منٹ
SDG: SDG 4 – معیاری تعلیم
سبق کا عنوان: نردبان اردو: دو، تین اور چار حرفی الفاظ (دہرائی اور مشق)

تعلیمی مقاصد
طلبہ دو، تین اور چار حرفی الفاظ کی پہچان کریں گے۔
طلبہ الفاظ درست تلفظ کے ساتھ پڑھیں اور لکھیں گے۔
طلبہ الفاظ کی درجہ بندی کر سکیں گے۔
حاصلِ تعلم / کامیابی کا معیار
طلبہ کم از کم 10 الفاظ درست پڑھیں گے۔
8 سے 10 الفاظ درست لکھیں گے۔
دو، تین اور چار حرفی الفاظ میں فرق کر سکیں گے۔
وقت	تعلیمی مواد و وسائل	تشکیلِ جائزہ	طلبہ کی سرگرمی / طریقۂ تدریس	استاد کی سرگرمی
5 منٹ	نردبان اردو، وائٹ بورڈ	زبانی سوالات	گزشتہ سبق کا اعادہ کریں۔	سبق کا مختصر تعارف دیں اور SDG 4 سے ربط قائم کریں۔
15 منٹ	درسی کتاب، فلیش کارڈز	زبانی جائزہ	دو، تین اور چار حرفی الفاظ پڑھیں۔	درست تلفظ کے ساتھ الفاظ پڑھوائیں اور رہنمائی کریں۔
15 منٹ	کاپی، وائٹ بورڈ	تحریری جائزہ	منتخب الفاظ لکھیں اور انہیں دو، تین اور چار حرفی خانوں میں تقسیم کریں۔	الفاظ لکھوائیں، تصحیح کریں اور انفرادی رہنمائی دیں۔
5 منٹ	وائٹ بورڈ	فیڈبیک	سبق کا خلاصہ بیان کریں اور چند الفاظ پڑھ کر سنائیں۔	سوالات پوچھیں، فیڈبیک دیں اور گھر کے لیے دہرائی کی ہدایت کریں۔
سبق کا جائزہ

طلبہ نے دو، تین اور چار حرفی الفاظ کی دہرائی کی، درست تلفظ کے ساتھ پڑھنے اور لکھنے کی مشق کی، اور الفاظ کو حروف کی تعداد کے مطابق درجہ بند کیا۔ سبق کے اختتام پر زبانی اور تحریری جائزے کے ذریعے طلبہ کی سیکھنے کی جانچ کی گئی اور SDG 4 (معیاری تعلیم) کی اہمیت کو اجاگر کیا گیا۔
"""

parsed = parse_lesson_paste(SAMPLE)
print("plan_title:", parsed.get("plan_title"))
print("class_value:", parsed.get("class_value"))
print("week:", parsed.get("week"))
print("club_period:", parsed.get("club_period"))
print("sdg:", parsed.get("sdg"))
print("lesson_title:", parsed.get("lesson_title"))
print("outcomes:", parsed.get("outcomes"))
print("success:", parsed.get("success_items"))
print("activities:", len(parsed.get("activities") or []))
for a in parsed.get("activities") or []:
    print(" ", a)
print("review:", (parsed.get("review_body") or "")[:80], "...")

assert parsed["class_value"] == "سوم"
assert "40 منٹ" in parsed["club_period"]
assert len(parsed["outcomes"]) == 3
assert len(parsed["success_items"]) == 3
assert len(parsed["activities"]) == 4
assert "نردبان" in parsed["lesson_title"]
assert "درجہ بند" in parsed["review_body"]

ROOT = Path(__file__).resolve().parent.parent
defaults = json.loads((ROOT / "app" / "default_values.json").read_text(encoding="utf-8"))
ctx = context_from_flat_fields({**defaults, **parsed})
out = ROOT / "outputs" / "paste-smoke.docx"
fill_template((ROOT / "templates" / "lesson_plan_template.docx").read_bytes(), ctx, out)

from docx import Document
doc = Document(str(out))
text = "\n".join(p.text for p in doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            text += "\n" + cell.text

checks = [
    "سنگل پیریڈ: 40 منٹ",
    "نردبان اردو",
    "دو، تین اور چار حرفی الفاظ کی پہچان",
    "کم از کم 10 الفاظ درست پڑھیں",
    "فلیش کارڈز",
    "درجہ بند کیا",
]
for c in checks:
    ok = c in text
    print(("OK" if ok else "MISSING"), c)
    assert ok

print("SMOKE_OK", out)
