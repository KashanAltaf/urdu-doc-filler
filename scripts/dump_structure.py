from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

path = Path("templates/Planner3_24August2026.docx")
doc = Document(str(path))
out = Path("templates/_structure.txt")
lines: list[str] = []


def cell_info(cell, prefix: str) -> None:
    lines.append(f"{prefix} text={cell.text!r}")
    for pi, p in enumerate(cell.paragraphs):
        align = p.alignment
        lines.append(f"{prefix}  p{pi} align={align} text={p.text!r}")
        for ri, r in enumerate(p.runs):
            lines.append(
                f"{prefix}    r{ri} font={r.font.name!r} size={r.font.size} bold={r.bold} text={r.text!r}"
            )


lines.append("=== PARAS ===")
for i, p in enumerate(doc.paragraphs):
    lines.append(f"P{i} style={p.style.name if p.style else None} text={p.text!r}")
    for ri, r in enumerate(p.runs):
        lines.append(f"  r{ri} font={r.font.name!r} size={r.font.size} text={r.text!r}")

for ti, table in enumerate(doc.tables):
    lines.append(f"=== TABLE {ti} {len(table.rows)}x{len(table.columns)} ===")
    # unique cell ids to detect merges
    for ri, row in enumerate(table.rows):
        lines.append(f"-- row {ri} --")
        seen = {}
        for ci, cell in enumerate(row.cells):
            cid = id(cell._tc)
            if cid in seen:
                lines.append(f"  C{ci}: MERGE_OF_C{seen[cid]}")
                continue
            seen[cid] = ci
            cell_info(cell, f"  C{ci}:")

out.write_text("\n".join(lines), encoding="utf-8")
print("wrote", out)
