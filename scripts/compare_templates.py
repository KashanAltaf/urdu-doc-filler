from pathlib import Path
from docx import Document
from zipfile import ZipFile
from lxml import etree

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
out = Path("templates/_compare_align.txt")
lines = []

def dump(path: Path, label: str) -> None:
    lines.append(f"===== {label} ({path.name}) =====")
    doc = Document(str(path))
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip() and i not in (0, 1, 5, 6):
            continue
        if p.text.strip() or i < 2:
            lines.append(f"P{i} align={p.alignment} text={p.text[:60]!r}")
    for ti, table in enumerate(doc.tables):
        lines.append(f"-- TABLE {ti} --")
        for ri, row in enumerate(table.rows[:3]):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs[:2]):
                    if p.text.strip() or (ri == 0 and pi == 0):
                        lines.append(
                            f"  T{ti}R{ri}C{ci}p{pi} align={p.alignment} text={p.text[:50]!r}"
                        )
    root = etree.fromstring(ZipFile(path).read("word/document.xml"))
    # sample first table first row jc/bidi/rtl counts
    bidi = len(root.findall(".//" + W + "bidi"))
    rtl = len(root.findall(".//" + W + "rtl"))
    lines.append(f"xml bidi={bidi} rtl={rtl}")
    tbl = root.findall(".//" + W + "tbl")[0]
    for i, tc in enumerate(tbl.findall(W + "tr")[0].findall(W + "tc")):
        p = tc.find(W + "p")
        if p is None:
            continue
        pPr = p.find(W + "pPr")
        jc = None
        bid = False
        if pPr is not None:
            j = pPr.find(W + "jc")
            jc = j.get(W + "val") if j is not None else None
            bid = pPr.find(W + "bidi") is not None
        text = "".join(t.text or "" for t in tc.findall(".//" + W + "t"))[:40]
        lines.append(f"  xml T0C{i} jc={jc} bidi={bid} text={text!r}")

dump(Path("templates/Planner3_24August2026.docx"), "PLANNER3")
dump(Path("templates/lesson formet in urdu 2026.docx"), "FORM2026")
out.write_text("\n".join(lines), encoding="utf-8")
print("wrote", out)
