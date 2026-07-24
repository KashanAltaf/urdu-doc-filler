from pathlib import Path
from docx import Document
from app.docx_fill import extract_placeholders

path = Path("templates/Planner3_24August2026.docx")
out = Path("templates/_inspect.txt")
lines: list[str] = []
lines.append(f"FILE: {path.name} SIZE: {path.stat().st_size}")
ph = extract_placeholders(path.read_bytes())
lines.append(f"PLACEHOLDERS: {ph if ph else '(none)'}")
doc = Document(str(path))
lines.append(f"PARAGRAPHS: {len(doc.paragraphs)}")
lines.append(f"TABLES: {len(doc.tables)}")
lines.append("--- BODY ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        lines.append(f"P{i}: {t}")
        fonts = sorted({(r.font.name or "") for r in p.runs})
        if fonts:
            lines.append(f"   fonts: {fonts}")
lines.append("--- TABLES ---")
for ti, table in enumerate(doc.tables):
    lines.append(f"TABLE {ti}: {len(table.rows)}r x {len(table.columns)}c")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip().replace("\n", " / ") for c in row.cells]
        lines.append(f"  R{ri}: " + " | ".join(cells))
lines.append("--- HEADERS/FOOTERS ---")
for si, section in enumerate(doc.sections):
    ht = "\n".join(p.text for p in section.header.paragraphs).strip()
    ft = "\n".join(p.text for p in section.footer.paragraphs).strip()
    if ht:
        lines.append(f"Section{si} header: {ht}")
    if ft:
        lines.append(f"Section{si} footer: {ft}")

# Also dump raw XML snippets that look like fillable blanks / underscores
xml = path.read_bytes()
# quick scan for common blank patterns in document.xml via python-docx parts
from zipfile import ZipFile

with ZipFile(path) as z:
    xml_text = z.read("word/document.xml").decode("utf-8")
underscores = xml_text.count("_")
lines.append(f"UNDERSCORE_CHARS_IN_XML: {underscores}")
out.write_text("\n".join(lines), encoding="utf-8")
print(f"wrote {out}")
