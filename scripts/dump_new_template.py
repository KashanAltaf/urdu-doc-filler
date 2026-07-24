from pathlib import Path
from docx import Document
from zipfile import ZipFile

path = Path("templates/lesson formet in urdu 2026.docx")
doc = Document(str(path))
out = Path("templates/_new_structure.txt")
lines: list[str] = []

lines.append(f"FILE: {path.name} SIZE: {path.stat().st_size}")
lines.append(f"PARAS: {len(doc.paragraphs)} TABLES: {len(doc.tables)}")

with ZipFile(path) as z:
    media = [n for n in z.namelist() if n.startswith("word/media/")]
    headers = [n for n in z.namelist() if "header" in n]
    lines.append(f"media: {media}")
    lines.append(f"headers: {headers}")

lines.append("=== PARAS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if t.strip() or p.runs:
        fonts = sorted({(r.font.name or "", str(r.font.size)) for r in p.runs})
        lines.append(f"P{i} align={p.alignment} style={p.style.name if p.style else None} text={t!r}")
        lines.append(f"   fonts={fonts}")
        for ri, r in enumerate(p.runs):
            has_d = "drawing" in r._element.xml or "pict" in r._element.xml
            if (r.text and r.text.strip()) or has_d:
                lines.append(f"   r{ri} size={r.font.size} drawing={has_d} text={r.text!r}")

lines.append("=== TABLES ===")
for ti, table in enumerate(doc.tables):
    lines.append(f"TABLE {ti}: {len(table.rows)}r x {len(table.columns)}c")
    for ri, row in enumerate(table.rows):
        seen = {}
        lines.append(f"-- row {ri} --")
        for ci, cell in enumerate(row.cells):
            cid = id(cell._tc)
            if cid in seen:
                lines.append(f"  C{ci}: MERGE_OF_C{seen[cid]}")
                continue
            seen[cid] = ci
            lines.append(f"  C{ci}: text={cell.text!r}")
            for pi, p in enumerate(cell.paragraphs):
                if not p.text.strip() and not p.runs:
                    continue
                sizes = [r.font.size for r in p.runs if (r.text or "").strip()]
                lines.append(f"    p{pi} align={p.alignment} sizes={sizes} text={p.text!r}")

out.write_text("\n".join(lines), encoding="utf-8")
print("wrote", out)
