"""Build a docxtpl template from the sample Planner3 document.

Preserves original paragraph alignment, font size, and run formatting
from Planner3_24August2026.docx when inserting placeholders.
"""

from __future__ import annotations

import json
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "templates" / "Planner3_24August2026.docx"
DST = ROOT / "templates" / "lesson_plan_template.docx"
DEFAULTS = ROOT / "app" / "default_values.json"
FONT = "Jameel Noori Nastaleeq"
MAX_LIST = 8
MAX_ACT = 8


def _format_from_paragraph(paragraph):
    """Clone rPr + size from the best existing text run (prefer non-empty text)."""
    candidates = []
    for run in paragraph.runs:
        el = run._element
        if el.find(qn("w:drawing")) is not None and not (run.text or "").strip():
            continue
        candidates.append(run)

    for run in candidates:
        if (run.text or "").strip():
            rPr = run._element.find(qn("w:rPr"))
            return (deepcopy(rPr) if rPr is not None else None), run.font.size

    for run in candidates:
        rPr = run._element.find(qn("w:rPr"))
        if rPr is not None:
            return deepcopy(rPr), run.font.size

    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None:
        rPr = pPr.find(qn("w:rPr"))
        if rPr is not None:
            return deepcopy(rPr), None
    return None, None


def set_para_text(paragraph, text: str) -> None:
    """Replace text runs but keep drawings, alignment, and original run formatting."""
    rPr_clone, size = _format_from_paragraph(paragraph)

    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:r"):
            continue
        if child.find(qn("w:drawing")) is not None or child.find(qn("w:pict")) is not None:
            continue
        p.remove(child)

    run = paragraph.add_run(text)
    if rPr_clone is not None:
        existing = run._element.find(qn("w:rPr"))
        if existing is not None:
            run._element.remove(existing)
        run._element.insert(0, rPr_clone)
    else:
        run.font.name = FONT
        if size is not None:
            run.font.size = size


def set_cell_lines(cell, lines: list[str]) -> None:
    """Fill a cell with lines, cloning the first paragraph's formatting for each line."""
    if not cell.paragraphs:
        return

    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)

    first = cell.paragraphs[0]
    if not lines:
        set_para_text(first, "")
        return

    set_para_text(first, lines[0])
    anchor = first._p
    for _ in lines[1:]:
        clone = deepcopy(anchor)
        anchor.addnext(clone)
        anchor = clone

    paras = cell.paragraphs
    for i, line in enumerate(lines):
        set_para_text(paras[i], line)


def delete_row(table, row_idx: int) -> None:
    table._tbl.remove(table.rows[row_idx]._tr)


def ensure_data_rows(table, count: int) -> None:
    """Ensure table has 1 header + `count` data rows."""
    while len(table.rows) > count + 1:
        delete_row(table, len(table.rows) - 1)
    while len(table.rows) < count + 1:
        table._tbl.append(deepcopy(table.rows[-1]._tr))


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"Missing source: {SRC}")

    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    outcomes = [
        "طلبہ نظم کے نئے الفاظ اور ان کے معانی سیکھیں گے۔",
        "طلبہ الفاظ کو استعمال کرتے ہوئے درست جملے بنائیں گے۔",
        "طلبہ نظم سے متعلق سوالات کے درست جوابات دے سکیں گے۔",
        "طلبہ نظم کے مرکزی خیال کو سمجھ سکیں گے۔",
    ]
    success_items = [
        "طلبہ کم از کم 10 الفاظ کے معانی درست بتائیں گے۔",
        "دیے گئے الفاظ سے درست جملے بنائیں گے۔",
        "تین سوالات کے درست جوابات لکھیں گے۔",
        "نظم کا مفہوم اپنے الفاظ میں بیان کریں گے۔",
    ]
    activities = [
        {
            "time": "10 منٹ",
            "materials": "درسی کتاب، وائٹ بورڈ",
            "assessment": "سوال و جواب",
            "student": "گزشتہ سبق کا اعادہ کریں اور نظم پڑھیں۔",
            "teacher": "سبق کا مختصر جائزہ اور SDG 4 کی اہمیت بیان کریں۔",
        },
        {
            "time": "20 منٹ",
            "materials": "درسی کتاب، چارٹ",
            "assessment": "زبانی سوالات",
            "student": "نئے الفاظ پڑھیں، معانی یاد کریں اور دہرائیں۔",
            "teacher": "الفاظ کے معانی مثالوں کے ساتھ سمجھائیں۔",
        },
        {
            "time": "20 منٹ",
            "materials": "کتاب، کاپی",
            "assessment": "انفرادی جائزہ",
            "student": "دیے گئے الفاظ سے جملے بنائیں۔",
            "teacher": "طلبہ کی رہنمائی کریں اور جملوں کی تصحیح کریں۔",
        },
        {
            "time": "20 منٹ",
            "materials": "درسی کتاب",
            "assessment": "تحریری جائزہ",
            "student": "نظم کے تین سوالات کے جوابات لکھیں۔",
            "teacher": "سوالات کی وضاحت کریں اور انفرادی مدد فراہم کریں۔",
        },
        {
            "time": "10 منٹ",
            "materials": "وائٹ بورڈ",
            "assessment": "زبانی فیڈبیک",
            "student": "سبق کا خلاصہ بیان کریں اور سوالات کے جواب دیں۔",
            "teacher": "سبق کا جائزہ لیں، فیڈبیک دیں اور گھر کے لیے دہرائی کی ہدایت کریں۔",
        },
    ]

    defaults = {
        "school_name": "دی لنکس اسکول",
        "plan_title": "سبق کا منصوبہ (جماعت سوم)",
        "subject_header": "سبق کا عنوان                             /موضوع",
        "lesson_title": "سبق کا عنوان: نظم: حمد (الفاظ/معانی، الفاظ سے جملے، سوال/جواب)",
        "sdg": "SDG: SDG 4 – معیاری تعلیم",
        "date": "۲۴ اگست",
        "date_label": "تاریخ",
        "class_value": "سوم",
        "class_label": "جماعت",
        "week": "دوسرا                 ہفتہ",
        "club_period": "کلب پیریڈ: 80 منٹ",
        "outcomes_label": "حاصلِ تعلم / کامیابی کے معیار",
        "objectives_label": "تعلیمی مقاصد",
        "col_time": "وقت",
        "col_materials": "تعلیمی مواد و وسائل",
        "col_assessment": "تشکیلِ جائزہ",
        "col_student": "طلبہ کی سرگرمی / طریقۂ تدریس",
        "col_teacher": "استاد کی سرگرمی",
        "review_heading": "سبق کا جائزہ",
        "review_body": (
            'طلبہ نے نظم "حمد" کے الفاظ اور معانی سیکھے، نئے الفاظ سے جملے بنائے، '
            "تین سوالات کے جوابات تحریر کیے اور نظم کے مرکزی خیال کو سمجھا۔ "
            "اس سبق کے ذریعے SDG 4 (معیاری تعلیم) کی اہمیت کو بھی اجاگر کیا گیا۔"
        ),
        "outcomes": outcomes,
        "success_items": success_items,
        "activities": activities,
    }
    for i in range(1, MAX_LIST + 1):
        defaults[f"outcome_{i}"] = outcomes[i - 1] if i <= len(outcomes) else ""
        defaults[f"success_{i}"] = success_items[i - 1] if i <= len(success_items) else ""
    for i in range(1, MAX_ACT + 1):
        if i <= len(activities):
            for key, value in activities[i - 1].items():
                defaults[f"a{i}_{key}"] = value
        else:
            for key in ("time", "materials", "assessment", "student", "teacher"):
                defaults[f"a{i}_{key}"] = ""

    # Preserve original sizes/alignment — inherit from Planner3 cells
    set_para_text(doc.paragraphs[0], "{{school_name}}")
    set_para_text(doc.paragraphs[1], "{{plan_title}}")
    set_para_text(doc.paragraphs[5], "{{review_heading}}: ")

    t0 = doc.tables[0]
    set_cell_lines(
        t0.rows[0].cells[0],
        ["{{subject_header}}", "{{lesson_title}}", "{{sdg}}"],
    )
    set_cell_lines(t0.rows[0].cells[1], ["{{date}}"])
    set_cell_lines(t0.rows[0].cells[2], ["{{date_label}}"])
    set_cell_lines(t0.rows[0].cells[3], ["{{class_value}}"])
    set_cell_lines(t0.rows[0].cells[4], ["{{class_label}}"])
    set_cell_lines(t0.rows[0].cells[5], ["{{week}}"])

    set_cell_lines(doc.tables[1].rows[0].cells[0], ["{{club_period}}"])

    t2 = doc.tables[2]
    set_cell_lines(
        t2.rows[0].cells[0],
        [f"{{{{success_{i}}}}}" for i in range(1, MAX_LIST + 1)],
    )
    set_cell_lines(t2.rows[0].cells[1], ["{{outcomes_label}}"])
    set_cell_lines(
        t2.rows[0].cells[2],
        [f"{{{{outcome_{i}}}}}" for i in range(1, MAX_LIST + 1)],
    )
    set_cell_lines(t2.rows[0].cells[3], ["{{objectives_label}}"])

    t3 = doc.tables[3]
    set_cell_lines(t3.rows[0].cells[0], ["{{col_time}}"])
    set_cell_lines(t3.rows[0].cells[1], ["{{col_materials}}"])
    set_cell_lines(t3.rows[0].cells[2], ["{{col_assessment}}"])
    set_cell_lines(t3.rows[0].cells[3], ["{{col_student}}"])
    set_cell_lines(t3.rows[0].cells[4], ["{{col_teacher}}"])

    ensure_data_rows(t3, MAX_ACT)
    for i in range(1, MAX_ACT + 1):
        cells = t3.rows[i].cells
        set_cell_lines(cells[0], [f"{{{{a{i}_time}}}}"])
        set_cell_lines(cells[1], [f"{{{{a{i}_materials}}}}"])
        set_cell_lines(cells[2], [f"{{{{a{i}_assessment}}}}"])
        set_cell_lines(cells[3], [f"{{{{a{i}_student}}}}"])
        set_cell_lines(cells[4], [f"{{{{a{i}_teacher}}}}"])

    set_cell_lines(doc.tables[4].rows[0].cells[0], ["{{review_body}}"])

    doc.save(str(DST))
    DEFAULTS.write_text(json.dumps(defaults, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {DST}")
    print(f"Wrote {DEFAULTS}")


if __name__ == "__main__":
    main()
